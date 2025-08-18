from fastapi import APIRouter, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from typing import Dict, Any, Optional
import requests
import io
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile
import os
from datetime import datetime
import re

router = APIRouter(prefix="/api/ct01", tags=["CT01"])

class CT01FormData(BaseModel):
    formData: Dict[str, Any]
    cccdData: Optional[Dict[str, Any]] = None
    template: Optional[Dict[str, Any]] = None
    type: str = "docx"

class CT01SubmitData(BaseModel):
    formData: Dict[str, Any]
    cccdData: Optional[Dict[str, Any]] = None

@router.post("/generate")
async def generate_ct01_file(data: CT01FormData):
    """
    Tạo file CT01 từ file DOCX gốc với data đã điền
    """
    try:
        print(f"Received data: {data.formData}")
        print(f"Template info: {data.template}")
        
        # Lấy file DOCX gốc từ Supabase
        docx_url = "https://rjrqtogyzmgyqvryxfyk.supabase.co/storage/v1/object/public/bieumau/ct01.docx"
        print(f"Downloading DOCX template from: {docx_url}")
        
        # Download file DOCX từ Supabase
        response = requests.get(docx_url)
        if response.status_code != 200:
            raise HTTPException(status_code=500, detail="Không thể tải file template từ Supabase")
        
        # Tạo file tạm để lưu DOCX
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
            temp_file.write(response.content)
            temp_file_path = temp_file.name
        
        try:
            # Điền data vào file DOCX
            filled_docx_bytes = fill_docx_template_with_data(temp_file_path, data.formData, data.cccdData)
            print(f"Data filled into DOCX template")
            
            if data.type == "docx":
                # Trả về DOCX
                filename = f"CT01-{data.template.get('code', 'CT01')}.docx"
                return StreamingResponse(
                    io.BytesIO(filled_docx_bytes),
                    media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    headers={"Content-Disposition": f"attachment; filename={filename}"}
                )
            elif data.type == "pdf":
                # Convert DOCX to PDF
                pdf_bytes = convert_docx_to_pdf(filled_docx_bytes)
                filename = f"CT01-{data.template.get('code', 'CT01')}.pdf"
                return StreamingResponse(
                    io.BytesIO(pdf_bytes),
                    media_type="application/pdf",
                    headers={"Content-Disposition": f"attachment; filename={filename}"}
                )
            else:
                # Trả về DOCX mặc định
                filename = f"CT01-{data.template.get('code', 'CT01')}.docx"
                return StreamingResponse(
                    io.BytesIO(filled_docx_bytes),
                    media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    headers={"Content-Disposition": f"attachment; filename={filename}"}
                )
        
        finally:
            # Xóa file tạm
            if os.path.exists(temp_file_path):
                os.unlink(temp_file_path)
        
    except Exception as e:
        print(f"Error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error generating file: {str(e)}")

@router.post("/submit")
async def submit_ct01_form(data: CT01SubmitData):
    """
    Nộp form CT01 trực tuyến
    """
    try:
        # TODO: Implement online submission logic
        # Có thể lưu vào database, gửi email, etc.
        
        return {
            "success": True,
            "message": "Form submitted successfully",
            "reference_id": f"CT01-{len(data.formData)}-{hash(str(data.formData)) % 10000}"
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error submitting form: {str(e)}")

def fill_docx_template_with_data(docx_path: str, form_data: Dict[str, Any], cccd_data: Dict[str, Any] = None) -> bytes:
    """
    Điền data vào file DOCX template
    """
    print(f"Filling DOCX template with data: {form_data}")
    if cccd_data:
        print(f"CCCD data available: {cccd_data}")
    
    # Debug: Log tất cả các trường có data
    print("📊 Available form fields:")
    for key, value in form_data.items():
        if value:  # Chỉ log những trường có giá trị
            print(f"   {key}: '{value}'")
    
    # Mở file DOCX
    doc = Document(docx_path)
    
    # Lấy ngày hiện tại cho phần chữ ký
    current_date = datetime.now()
    formatted_date = f"ngày {current_date.day} tháng {current_date.month} năm {current_date.year}"
    
    # Helper function to get data with fallback to CCCD
    def get_field_value(field_name: str, cccd_field: str = None) -> str:
        # Ưu tiên form_data, fallback sang cccd_data nếu có
        value = form_data.get(field_name, "")
        if not value and cccd_data and cccd_field:
            value = cccd_data.get(cccd_field, "")
        return str(value) if value else ""
    
    # Hàm tách ngày, tháng, năm từ chuỗi "DD/MM/YYYY"
    def parse_date(date_str):
        try:
            if "-" in date_str:  # yyyy-mm-dd format
                parts = date_str.split("-")
                return parts[2], parts[1], parts[0]
            elif "/" in date_str:  # dd/mm/yyyy format
                parts = date_str.split("/")
                return parts[0], parts[1], parts[2] if len(parts) == 3 else ""
            else:
                return "", "", ""
        except ValueError:
            return "", "", ""
    
    # Hàm tách số định danh thành 12 ô
    def split_id_number(id_number):
        id_str = str(id_number).replace(" ", "").replace("-", "")
        return list(id_str) if len(id_str) == 12 else [''] * 12
    
    def clean_dots_and_fill(text, value):
        """Xóa tất cả dấu chấm và thêm giá trị vào cuối"""
        if not value:
            return text
                
        cleaned_text = re.sub(r'\.+', '', text)  
        cleaned_text = re.sub(r'\.', '', cleaned_text) 
        
        # Làm sạch khoảng trắng thừa
        cleaned_text = re.sub(r'\s+', ' ', cleaned_text)  # Thay thế nhiều khoảng trắng bằng 1 khoảng trắng
        cleaned_text = cleaned_text.strip()
        
        # Thêm giá trị
        if cleaned_text.endswith(":"):
            return f"{cleaned_text} {value}"
        else:
            return f"{cleaned_text}: {value}"
    
    def apply_font_formatting(paragraph, text):
        """Áp dụng font Times New Roman, cỡ 13 cho text"""
        paragraph.text = text
        for run in paragraph.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(13)
    
    def apply_font_to_cell(cell, text):
        """Áp dụng font Times New Roman, cỡ 13 cho ô trong bảng"""
        cell.text = text
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(13)


    
    # Debug: In ra tất cả paragraphs để xem cấu trúc template
    print("Debug: Tìm kiếm các trường trong template...")
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip():
            print(f"Paragraph {i}: '{paragraph.text}'")
            # Kiểm tra các pattern cụ thể
            if "Kính gửi" in paragraph.text:
                print(f"  -> Found 'Kính gửi' pattern")
            if "Họ, chữ đệm và tên" in paragraph.text:
                print(f"  -> Found 'Họ tên' pattern")
            if "Ngày, tháng, năm sinh" in paragraph.text:
                print(f"  -> Found 'Ngày sinh' pattern")
            if "Số điện thoại" in paragraph.text:
                print(f"  -> Found 'SĐT' pattern")
            if "chủ hộ" in paragraph.text:
                print(f"  -> Found 'Chủ hộ' pattern")
            if "Nội dung đề nghị" in paragraph.text:
                print(f"  -> Found 'Nội dung đề nghị' pattern")
    
    # Thay thế thông tin trong các đoạn văn
    replacements_made = 0
    
    # Lấy nội dung đề nghị trước để sử dụng cho các trường khác
    noi_dung = get_field_value("noi_dung_de_nghi")
    
    # Lưu index của paragraph chứa "Nội dung đề nghị"
    noi_dung_paragraph_idx = -1
    
    for i, paragraph in enumerate(doc.paragraphs):
        original_text = paragraph.text
        
        # Kính gửi - pattern đơn giản
        if "Kính gửi(1):" in original_text:
            co_quan_tiep_nhan = get_field_value("co_quan_tiep_nhan")
            if co_quan_tiep_nhan:
                apply_font_formatting(paragraph, f"Kính gửi(1): {co_quan_tiep_nhan}")
                replacements_made += 1
                print(f"Filled 'Kính gửi': {co_quan_tiep_nhan}")
        
        # Họ, chữ đệm và tên - pattern đơn giản
        if "1. Họ, chữ đệm và tên:" in original_text:
            ho_ten = get_field_value("ho_ten", "personName")
            if ho_ten:
                apply_font_formatting(paragraph, f"1. Họ, chữ đệm và tên: {ho_ten}")
                replacements_made += 1
                print(f"Filled 'Họ tên': {ho_ten}")
        
        # Ngày sinh và giới tính - pattern đơn giản
        if "2. Ngày, tháng, năm sinh:" in original_text:
            ngay_sinh = get_field_value("ngay_sinh", "dateOfBirth")
            gioi_tinh = get_field_value("gioi_tinh", "gender")
            if ngay_sinh:
                day, month, year = parse_date(ngay_sinh)
                date_str = f"{day} / {month} / {year}"
                apply_font_formatting(paragraph, f"2. Ngày, tháng, năm sinh: {date_str}       3. Giới tính: {gioi_tinh}")
                replacements_made += 1
                print(f"Filled 'Ngày sinh': {date_str}, Giới tính: {gioi_tinh}")
        
        # Số điện thoại và Email - pattern đơn giản
        if "5. Số điện thoại liên hệ:" in original_text:
            sdt = get_field_value("dien_thoai") or get_field_value("so_dien_thoai")
            email = get_field_value("email")
            if sdt:
                apply_font_formatting(paragraph, f"5. Số điện thoại liên hệ: {sdt}    6. Email: {email}")
                replacements_made += 1
                print(f"Filled 'SĐT': {sdt}, Email: {email}")
        
        # Chủ hộ và mối quan hệ - pattern đơn giản
        if "7. Họ, chữ đệm và tên chủ hộ:" in original_text:
            chu_ho = get_field_value("chu_ho") or get_field_value("ho_ten_chu_ho")
            quan_he = get_field_value("quan_he_chu_ho") or get_field_value("moi_quan_he_chu_ho", "") or "Chủ hộ"
            if chu_ho:
                apply_font_formatting(paragraph, f"7. Họ, chữ đệm và tên chủ hộ: {chu_ho}    8. Mối quan hệ với chủ hộ: {quan_he}")
                replacements_made += 1
                print(f"Filled 'Chủ hộ': {chu_ho}, Quan hệ: {quan_he}")
        
        # Đánh dấu paragraph chứa "Nội dung đề nghị" - pattern mở rộng
        if "10. Nội dung đề nghị" in original_text:
            noi_dung_paragraph_idx = i
            print(f"🎯 Found 'Nội dung đề nghị' at paragraph {i}: '{original_text}'")
    
    # Xử lý "Nội dung đề nghị" sau khi tìm được paragraph
    if noi_dung_paragraph_idx != -1 and noi_dung:
        import re
        
        # Tìm tất cả các paragraph liên quan (bao gồm cả dòng chấm)
        paragraphs_to_process = [noi_dung_paragraph_idx]
        
        # Tìm các paragraph tiếp theo chỉ chứa dấu chấm hoặc trống
        for idx in range(noi_dung_paragraph_idx + 1, min(noi_dung_paragraph_idx + 10, len(doc.paragraphs))):
            if idx >= len(doc.paragraphs):
                break
                
            next_paragraph = doc.paragraphs[idx]
            next_text = next_paragraph.text.strip()
            
            # Kiểm tra xem có phải dòng chấm không
            meaningful_content = re.sub(r'[.…\-_\s]', '', next_text)
            if (next_text == '' or 
                len(meaningful_content) == 0 or
                re.match(r'^[ .…\-_\s]*$', next_text)):
                
                paragraphs_to_process.append(idx)
                print(f"✓ Added to processing list: paragraph {idx}: '{next_text}'")
            else:
                print(f"✗ Found content, stopping: paragraph {idx}: '{next_text}'")
                break
        
        # Xử lý paragraph đầu tiên (chứa tiêu đề)
        main_paragraph = doc.paragraphs[noi_dung_paragraph_idx]
        original_text = main_paragraph.text
        
        # Làm sạch text gốc
        cleaned_text = re.sub(r'[.…\-_]+', '', original_text)
        cleaned_text = re.sub(r'\s+', ' ', cleaned_text).strip()
        if not cleaned_text.endswith(':'):
            cleaned_text += ':'
        
        # Xóa nội dung paragraph chính và điền lại
        main_paragraph.clear()
        run = main_paragraph.add_run(f"{cleaned_text} {noi_dung}")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(13)
        
        # Xóa hoặc làm rỗng các paragraph chấm thừa
        removed_count = 0
        for idx in paragraphs_to_process[1:]:  # Bỏ qua paragraph đầu tiên
            if idx < len(doc.paragraphs):
                try:
                    p = doc.paragraphs[idx]
                    # Thử method 1: Clear và set text rỗng
                    p.clear()
                    p.text = ""
                    
                    # Thử method 2: Xóa hoàn toàn nếu có thể
                    try:
                        p_element = p._element
                        parent = p_element.getparent()
                        if parent is not None:
                            parent.remove(p_element)
                            print(f"✅ Completely removed paragraph at index {idx}")
                        else:
                            print(f"✅ Cleared content of paragraph at index {idx}")
                    except:
                        print(f"✅ Cleared content of paragraph at index {idx}")
                    
                    removed_count += 1
                except Exception as e:
                    print(f"❌ Warning: Could not process paragraph at index {idx}: {e}")
        
        replacements_made += 1
        print(f"✅ Filled 'Nội dung đề nghị': {noi_dung}")
        print(f"📊 Processed {removed_count} empty/dotted paragraphs")
        print(f"🔍 Original: '{original_text}'")
        print(f"🔍 Cleaned: '{cleaned_text} {noi_dung}'")
    
    print(f"Total paragraph replacements made: {replacements_made}")
    
    # Xử lý các bảng
    print(f"🔍 Processing {len(doc.tables)} tables...")
    for table_idx, table in enumerate(doc.tables):
        print(f"📋 Table {table_idx}: Rows: {len(table.rows)}, Columns: {len(table.rows[0].cells) if table.rows else 0}")
        
        if table.rows and table.rows[0].cells:
            first_cell_text = table.rows[0].cells[0].text.strip()
            print(f"   First cell text: '{first_cell_text}'")
            
            # Bảng "4. Số định danh cá nhân"
            if "4. Số định danh cá nhân:" in first_cell_text:
                so_dinh_danh = get_field_value("so_dinh_danh", "idCode") or get_field_value("so_cccd", "idCode")
                print(f"✅ Found ID table. ID: '{so_dinh_danh}', Cells: {len(table.rows[0].cells)}")
                if so_dinh_danh and len(table.rows[0].cells) >= 13:
                    id_digits = split_id_number(so_dinh_danh)
                    print(f"🔢 ID digits: {id_digits}")
                    for i in range(12):
                        if i + 1 < len(table.rows[0].cells):
                            apply_font_to_cell(table.rows[0].cells[i + 1], id_digits[i])
                            print(f"   Cell {i+1}: '{id_digits[i]}'")
                else:
                    print(f"❌ Cannot fill ID: so_dinh_danh='{so_dinh_danh}', cells={len(table.rows[0].cells)}")
            
            # Bảng "9. Số định danh cá nhân của chủ hộ"
            elif "9. Số định danh cá nhân của chủ hộ:" in first_cell_text:
                so_dinh_danh_chu_ho = get_field_value("dinh_danh_chu_ho") or get_field_value("so_dinh_danh_chu_ho")
                print(f"✅ Found Head ID table. Head ID: '{so_dinh_danh_chu_ho}'")
                if so_dinh_danh_chu_ho and len(table.rows[0].cells) >= 13:
                    id_digits = split_id_number(so_dinh_danh_chu_ho)
                    print(f"🔢 Head ID digits: {id_digits}")
                    for i in range(12):
                        if i + 1 < len(table.rows[0].cells):
                            apply_font_to_cell(table.rows[0].cells[i + 1], id_digits[i])
                else:
                    print(f"❌ Cannot fill Head ID: so_dinh_danh_chu_ho='{so_dinh_danh_chu_ho}'")
        
            # Bảng "11. Những thành viên trong hộ gia đình cùng thay đổi"
            elif first_cell_text == "TT":
                print(f"✅ Found family members table")
                # Xóa các hàng hiện có (trừ hàng tiêu đề)
                original_rows = len(table.rows)
                while len(table.rows) > 1:
                    table._element.remove(table.rows[-1]._element)
                print(f"🗑️ Removed {original_rows - 1} existing rows")
                
                # Thêm hàng mới cho từng thành viên
                members = form_data.get("thanh_vien_gia_dinh") or form_data.get("thanh_vien_ho_gia_dinh") or form_data.get("thanh_vien_cung_thay_doi")
                print(f"👥 Found {len(members) if members else 0} family members to add")
                if members and isinstance(members, list):
                    for i, member in enumerate(members, 1):
                        row = table.add_row()
                        if len(row.cells) >= 6:
                            apply_font_to_cell(row.cells[0], str(i))  # Số thứ tự
                            apply_font_to_cell(row.cells[1], member.get('ho_ten', ''))
                            apply_font_to_cell(row.cells[2], member.get('ngay_sinh', ''))
                            apply_font_to_cell(row.cells[3], member.get('gioi_tinh', ''))
                            apply_font_to_cell(row.cells[4], member.get('so_dinh_danh', ''))
                            apply_font_to_cell(row.cells[5], member.get('moi_quan_he', member.get('quan_he', '')))
                            print(f"   Row {i}: {member.get('ho_ten', '')}")
                else:
                    print(f"ℹ️ No family members data to fill")
            
            # Bảng chữ ký chủ sở hữu
            elif "Họ và tên chủ sở hữu" in first_cell_text or "chủ sở hữu" in first_cell_text.lower():
                print(f"✅ Found owner signature table")
                chu_so_huu = get_field_value("chu_so_huu_ho_ten")
                chu_so_huu_id = get_field_value("chu_so_huu_dinh_danh")
                if chu_so_huu:
                    # Tìm và điền thông tin chủ sở hữu
                    for row in table.rows:
                        for cell_idx, cell in enumerate(row.cells):
                            if "Họ và tên:" in cell.text:
                                cell.text = f"Họ và tên: {chu_so_huu}"
                                apply_font_to_cell(cell, f"Họ và tên: {chu_so_huu}")
                            elif "Số định danh cá nhân:" in cell.text and chu_so_huu_id:
                                cell.text = f"Số định danh cá nhân: {chu_so_huu_id}"
                                apply_font_to_cell(cell, f"Số định danh cá nhân: {chu_so_huu_id}")
                    print(f"📝 Filled owner info: {chu_so_huu}")
            
            # Bảng chữ ký cha/mẹ/người giám hộ
            elif "cha/mẹ" in first_cell_text.lower() or "giám hộ" in first_cell_text.lower():
                print(f"✅ Found guardian signature table")
                giam_ho = get_field_value("giam_ho_ho_ten")
                giam_ho_id = get_field_value("giam_ho_dinh_danh")
                if giam_ho:
                    # Tìm và điền thông tin người giám hộ
                    for row in table.rows:
                        for cell in row.cells:
                            if "Họ và tên:" in cell.text:
                                cell.text = f"Họ và tên: {giam_ho}"
                                apply_font_to_cell(cell, f"Họ và tên: {giam_ho}")
                            elif "Số định danh cá nhân:" in cell.text and giam_ho_id:
                                cell.text = f"Số định danh cá nhân: {giam_ho_id}"
                                apply_font_to_cell(cell, f"Số định danh cá nhân: {giam_ho_id}")
                    print(f"📝 Filled guardian info: {giam_ho}")
            
            # Bảng ý kiến và các bảng cuối - BỎ QUA, không xử lý
            elif "Ý KIẾN CỦA CHỦ HỘ" in first_cell_text or "Ý KIẾN CỦA CHA" in first_cell_text or "YÊU CẦU" in first_cell_text:
                print(f"⏭️ Skipped opinion/signature table: '{first_cell_text[:50]}...' (removed by request)")
            else:
                print(f"❓ Unknown table: '{first_cell_text[:50]}...'")
                # Debug: In ra tất cả các cell của bảng không nhận diện được
                if len(table.rows) > 0:
                    print(f"   Debug - Table structure:")
                    for row_idx, row in enumerate(table.rows[:3]):  # Chỉ xem 3 hàng đầu
                        for cell_idx, cell in enumerate(row.cells[:4]):  # Chỉ xem 4 cột đầu
                            cell_text = cell.text.strip()[:30]  # Chỉ lấy 30 ký tự đầu
                            print(f"     Row {row_idx}, Cell {cell_idx}: '{cell_text}'")
        else:
            print(f"❌ Table {table_idx}: No rows or cells")
    
    # Lưu file vào bytes
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output.getvalue()

def convert_docx_to_pdf(docx_bytes: bytes) -> bytes:
    """
    Convert DOCX to PDF
    """
    try:
        # Sử dụng python-docx2pdf hoặc weasyprint để convert
        # Tạm thời trả về DOCX bytes nếu không có converter
        return docx_bytes
    except Exception as e:
        print(f"PDF conversion failed: {e}")
        return docx_bytes

@router.post("/preview")
async def preview_ct01_document(request: dict):
    """Endpoint để tạo file DOCX cho preview"""
    try:
        form_data = request.get("formData", {})
        cccd_data = request.get("cccdData", {})
        
        # Tải template từ Supabase
        docx_url = "https://rjrqtogyzmgyqvryxfyk.supabase.co/storage/v1/object/public/bieumau/ct01.docx"
        response = requests.get(docx_url)
        if response.status_code != 200:
            raise HTTPException(status_code=500, detail="Không thể tải file template từ Supabase")
        
        # Tạo file tạm để lưu DOCX
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
            temp_file.write(response.content)
            temp_file_path = temp_file.name
        
        try:
            # Tạo file DOCX đã điền sẵn
            docx_content = fill_docx_template_with_data(temp_file_path, form_data, cccd_data)
        finally:
            # Xóa file tạm
            if os.path.exists(temp_file_path):
                os.unlink(temp_file_path)
        
        # Lưu file tạm thời cho preview
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
            tmp_file.write(docx_content)
            tmp_path = tmp_file.name
        
        # Trả về đường dẫn file tương đối
        file_id = os.path.basename(tmp_path)
        return {"file_path": file_id, "filename": "CT01-preview.docx"}
        
    except Exception as e:
        print(f"❌ Error in preview_ct01_document: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/preview/{file_id}")
async def serve_preview_file(file_id: str):
    """Serve file DOCX cho preview"""
    try:
        import os
        import tempfile
        
        # Tìm file trong thư mục temp
        temp_dir = tempfile.gettempdir()
        full_path = os.path.join(temp_dir, file_id)
        
        if not os.path.exists(full_path):
            raise HTTPException(status_code=404, detail="File not found")
        
        with open(full_path, "rb") as f:
            content = f.read()
        
        return StreamingResponse(
            io.BytesIO(content),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "inline; filename=CT01-preview.docx"}
        )
        
    except Exception as e:
        print(f"❌ Error serving preview file: {e}")
        raise HTTPException(status_code=500, detail=str(e))