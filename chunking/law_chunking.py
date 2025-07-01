import re
import json
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
import os
import glob
import re
import json
from docx import Document
import os
import glob

def read_docx_all(filepath):
    from docx import Document
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    doc = Document(filepath)
    texts = []
    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            texts.append(block.text)
        elif isinstance(block, Table):
            for row in block.rows:
                for cell in row.cells:
                    texts.append(cell.text)
    # Gộp lại thành dòng, tránh lặp và dòng rỗng
    lines = []
    for line in texts:
        for l2 in line.split('\n'):
            l2 = l2.strip()
            if l2 and l2 not in lines:
                lines.append(l2)
    return "\n".join(lines)

# Hàm duyệt block item (paragraph/table) đúng theo thứ tự gốc file
def iter_block_items(parent):
    """
    Yield paragraphs and tables exactly as in the Word file, giữ đúng thứ tự thật.
    """
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    for child in parent.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)



def extract_law_meta(raw_text: str):
    lines = [l.strip() for l in raw_text.split("\n") if l.strip()]

    # Số hiệu - tìm trong tối đa 20 dòng đầu và cả trong bảng
    law_code = None
    for i, line in enumerate(lines[:20]):
        # Không lấy số hiệu căn cứ
        if re.search(r"(Căn cứ|Nghị định|Điều|Quy định)", line, re.IGNORECASE):
            continue
        m = re.search(r"(Luật số|Số)[:：]?\s*([0-9]{1,4}/[0-9]{4}/?[A-Z0-9\-]+)", line, re.IGNORECASE)
        if m:
            print(f"[MATCH] Dòng {i}: {line} --> {m.group(2)}")
            law_code = m.group(2).strip(" .")
            break
        if re.match(r"^(Luật số|Số)[:：]?\s*$", line, re.IGNORECASE) and i + 1 < len(lines):
            next_line = lines[i + 1].strip()
            if re.search(r"(Căn cứ|Nghị định|Điều|Quy định)", next_line, re.IGNORECASE):
                continue
            m2 = re.match(r"^([0-9]{1,4}/[0-9]{4}/?[A-Z0-9\-]+)", next_line)
            if m2:
                print(f"[MATCH SPLIT] Dòng {i} + {i+1}: {next_line} --> {m2.group(1)}")
                law_code = m2.group(1)
                break
    if not law_code:
        law_code = "Không xác định"

    # law_name: tìm dòng đầu tiên bắt đầu bằng 'LUẬT', 'NGHỊ ĐỊNH', ... nhưng KHÔNG chứa 'số'
    law_name = None
    for i, line in enumerate(lines[:40]):
        if re.match(r"^(LUẬT|NGHỊ ĐỊNH|THÔNG TƯ|QUYẾT ĐỊNH)$", line, re.IGNORECASE):
            if i + 1 < len(lines):
                next_line = lines[i+1].strip()
                law_name = f"{line.title()} {next_line}"
                break
        # Nếu là 'LUẬT ...' hoặc 'NGHỊ ĐỊNH ...' nhưng KHÔNG chứa 'số' hoặc 'Số'
        m = re.match(r"^(LUẬT|NGHỊ ĐỊNH|THÔNG TƯ|QUYẾT ĐỊNH)\s+(.+)$", line, re.IGNORECASE)
        if m and not re.search(r"\bsố\b", line, re.IGNORECASE):
            law_name = line.title()
            break
    # Nếu vẫn chưa có, trả về 'Không xác định'
    if not law_name:
        law_name = "Không xác định"

    # promulgation_date
    promulgation_date = None
    for line in lines[:30]:
        m = re.search(r"ngày\s+(\d{1,2}\s+tháng\s+\d{1,2}\s+năm\s+\d{4})", line, re.IGNORECASE)
        if m:
            promulgation_date = m.group(1)
            break

    # promulgator: lấy dòng đầu tiên không rỗng
    promulgator = lines[0] if lines else None

    # law_type
    law_type = None
    if law_name and law_name != "Không xác định":
        prefix = law_name.split()[0].upper()
        if prefix == "LUẬT":
            law_type = "Luật"
        elif prefix == "NGHỊ":
            law_type = "Nghị định"
        elif prefix == "QUYẾT":
            law_type = "Quyết định"
        elif prefix == "THÔNG":
            law_type = "Thông tư"

    return {
        "law_name": law_name,
        "law_code": law_code or "Không xác định",
        "promulgation_date": promulgation_date,
        "effective_date": None,
        "promulgator": promulgator,
        "law_type": law_type
    }


# ---- PARSER PHẦN THÂN ----
def parse_law_text(raw, meta):
    # Pattern để lấy cả tiêu đề chương và tên chương
    chapter_pat = re.compile(r"(Chương [IVXLCDM]+\.*[^\n]*\n[^\n]*)", re.IGNORECASE)
    # Chỉ lấy khi có format chính xác "Điều {số}. {nội dung}"
    article_pat = re.compile(r"Điều\s+(\d+)\s*\.\s+([^\n]+)", re.IGNORECASE)
    clause_pat = re.compile(r"(^\d+\.)", re.MULTILINE)
    point_pat = re.compile(r"(^[a-e]\))", re.MULTILINE)

    chapters = [(m.start(), m.group()) for m in chapter_pat.finditer(raw)]
    if not chapters:
        chapters = [(0, "Văn bản không có chương")]
        chapters.append((len(raw), ""))
    else:
        chapters.append((len(raw), ""))

    records = []
    for i in range(len(chapters)-1):
        chap_start, chap_full_title = chapters[i]
        chap_end = chapters[i+1][0]
        chap_text = raw[chap_start:chap_end]
        
        # Tách tiêu đề chương và tên chương
        chap_lines = chap_full_title.strip().split('\n')
        if len(chap_lines) >= 2:
            chapter_title = chap_lines[0].strip()  # "Chương I"
            chapter_name = chap_lines[1].strip()   # "Tên chương"
            chapter_content = chapter_name  # Chỉ lấy tên chương, bỏ "Chương I"
        else:
            chapter_title = chap_full_title.strip()
            chapter_name = ""
            chapter_content = chapter_title
            
        chapter_id = f"{meta['law_code']}_{chapter_title.replace(' ','_').replace('.','')}"
        
        # Tìm articles với format chính xác "Điều {số}. {nội dung}"
        articles = []
        for match in article_pat.finditer(chap_text):
            article_number = match.group(1)
            article_content = match.group(2)
            article_title = f"Điều {article_number}"
            # Lưu cả vị trí bắt đầu và nội dung đầy đủ
            articles.append((match.start(), article_title, match.group(0)))
        
        # Luôn tạo record cho chương trước
        chapter_record = {
            **meta,
            "chapter": chapter_title,
            "section": None,
            "article": None,
            "clause": None,
            "point": None,
            "type": "chương",
            "id": chapter_id,
            "parent_id": None,
            "parent_type": None,
            "content": chapter_content,
            "law_ref": f"{meta['law_name']}, {chapter_title}",
            "category": "law"
        }
        records.append(chapter_record)
        
        if not articles:
            continue
        articles.append((len(chap_text), "", ""))
        for j in range(len(articles)-1):
            art_start, art_title, art_full_text = articles[j]
            art_end = articles[j+1][0]
            art_text = chap_text[art_start:art_end]
            article_id = f"{chapter_id}_D{art_title.split()[1]}"
            
            # Tìm điểm kết thúc thực sự của điều (trước phần phụ lục)
            article_content_end = len(art_text)
            
            # Tìm các marker kết thúc điều
            end_markers = [
                "Nơi nhận:",
                "PHỤ LỤC",
                "Mẫu số",
                "TM. CHÍNH PHỦ",
                "KT. THỦ TƯỚNG",
                "PHÓ THỦ TƯỚNG"
            ]
            
            for marker in end_markers:
                marker_pos = art_text.find(marker)
                if marker_pos != -1 and marker_pos < article_content_end:
                    article_content_end = marker_pos
            
            # Lấy nội dung điều (không bao gồm phần phụ lục)
            full_article_content = art_text[:article_content_end].strip()
            
            # Tìm khoản đầu tiên để cắt nội dung điều
            first_clause_match = clause_pat.search(full_article_content)
            if first_clause_match:
                # Nội dung điều chỉ là phần từ đầu đến trước khoản đầu tiên
                article_title_content = full_article_content[:first_clause_match.start()].strip()
            else:
                # Nếu không có khoản, lấy toàn bộ nội dung
                article_title_content = full_article_content
            
            # Bỏ "Điều X." khỏi content
            article_title_content = re.sub(r'^Điều\s+\d+\s*\.\s*', '', article_title_content).strip()
            
            # Luôn tạo record cho điều trước
            article_record = {
                **meta,
                "chapter": chapter_title,
                "section": None,
                "article": art_title,
                "clause": None,
                "point": None,
                "type": "điều",
                "id": article_id,
                "parent_id": chapter_id,
                "parent_type": "chương",
                "content": article_title_content,
                "law_ref": f"{meta['law_name']}, {art_title}",
                "category": "law"
            }
            records.append(article_record)
            
            # Xử lý khoản con chỉ trong nội dung điều
            clause_matches = list(clause_pat.finditer(full_article_content))
            if clause_matches:
                clause_spans = [m.start() for m in clause_matches] + [len(full_article_content)]
                for k in range(len(clause_spans)-1):
                    clause_start = clause_spans[k]
                    clause_end = clause_spans[k+1]
                    clause_text = full_article_content[clause_start:clause_end].strip()
                    clause_no = re.match(r"^(\d+)\.", clause_text)
                    clause_id = f"{article_id}_K{clause_no.group(1)}" if clause_no else f"{article_id}_K?"
                    
                    # Bỏ "1." khỏi content của khoản
                    if clause_no:
                        clause_content = clause_text[clause_no.end():].strip()
                    else:
                        clause_content = clause_text
                    point_matches = list(point_pat.finditer(clause_text))
                    if point_matches:
                        point_spans = [m.start() for m in point_matches] + [len(clause_text)]
                        for m_ in range(len(point_spans)-1):
                            point_start = point_spans[m_]
                            point_end = point_spans[m_+1]
                            point_text = clause_text[point_start:point_end].strip()
                            point_id_match = re.match(r"([a-e])\)", point_text)
                            point_id = f"{clause_id}_{point_id_match.group(1)}" if point_id_match else f"{clause_id}_?"
                            
                            # Bỏ "a)" khỏi content của điểm
                            if point_id_match:
                                point_content = point_text[point_id_match.end():].strip()
                            else:
                                point_content = point_text
                            records.append({
                                **meta,
                                "chapter": chapter_title,
                                "section": None,
                                "article": art_title,
                                "clause": clause_no.group(1) if clause_no else None,
                                "point": point_id_match.group(1) if point_id_match else None,
                                "type": "điểm",
                                "id": point_id,
                                "parent_id": clause_id,
                                "parent_type": "khoản",
                                "content": point_content,
                                "law_ref": f"{meta['law_name']}, {art_title}, Khoản {clause_no.group(1) if clause_no else '?'}{', Điểm ' + point_id_match.group(1) if point_id_match else ''}",
                                "category": "law"
                            })
                    else:
                        records.append({
                            **meta,
                            "chapter": chapter_title,
                            "section": None,
                            "article": art_title,
                            "clause": clause_no.group(1) if clause_no else None,
                            "point": None,
                            "type": "khoản",
                            "id": clause_id,
                            "parent_id": article_id,
                            "parent_type": "điều",
                            "content": clause_content,
                            "law_ref": f"{meta['law_name']}, {art_title}, Khoản {clause_no.group(1) if clause_no else '?'}",
                            "category": "law"
                        })
    return records


# ==== MAIN: XỬ LÝ FOLDER ====
FOLDER = "./laws"  # Đổi thành thư mục chứa các file .docx
os.makedirs("output_json", exist_ok=True)

docx_files = glob.glob(os.path.join(FOLDER, "*.docx"))
print(f"Tìm thấy {len(docx_files)} file .docx trong {FOLDER}")

# Gộp tất cả records vào 1 list
all_records = []

for file_path in docx_files:
    print(f"Đang xử lý: {file_path}")
    try:
        def read_docx(filepath):
            doc = Document(filepath)
            text = "\n".join([p.text for p in doc.paragraphs])
            return text
        raw = read_docx_all(file_path)
        meta = extract_law_meta(raw)
        print("META:", meta)
        records = parse_law_text(raw, meta)
        all_records.extend(records)
        print(f"Đã thêm {len(records)} records từ {file_path}")
    except Exception as e:
        print(f"[WARNING] Không thể xử lý file {file_path}: {e}")
        continue

# Xuất tất cả vào 1 file JSON duy nhất
out_path = os.path.join("chunking/output_json", "all_laws.json")
with open(out_path, "w", encoding="utf-8") as out:
    json.dump(all_records, out, ensure_ascii=False, indent=2)

print(f"Tổng số record: {len(all_records)} | Xuất file: {out_path}")
