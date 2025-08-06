#!/usr/bin/env python3
"""
Script để tạo file CT01 template với placeholder
"""
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_ct01_template():
    # Tạo document mới
    doc = Document()
    
    # Header
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run = header.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM")
    header_run.bold = True
    header_run.font.size = Pt(14)
    
    doc.add_paragraph("Độc lập - Tự do - Hạnh phúc", style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("* * *", style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Tiêu đề
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("TỜ KHAI THAY ĐỔI THÔNG TIN CƯ TRÚ")
    title_run.bold = True
    title_run.font.size = Pt(16)
    
    doc.add_paragraph("Biểu mẫu dùng cho công dân Việt Nam thay đổi thông tin cư trú.", style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Mẫu số: CT01", style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Kính gửi
    doc.add_paragraph()
    doc.add_paragraph("Kính gửi: {{CO_QUAN_TIEP_NHAN}}")
    
    # I. THÔNG TIN NGƯỜI KÊ KHAI
    doc.add_paragraph()
    section1 = doc.add_paragraph()
    section1_run = section1.add_run("I. THÔNG TIN NGƯỜI KÊ KHAI")
    section1_run.bold = True
    
    doc.add_paragraph("1. Họ, chữ đệm và tên: {{HO_TEN}}")
    doc.add_paragraph("2. Ngày, tháng, năm sinh: {{NGAY_SINH}}")
    doc.add_paragraph("3. Giới tính: {{GIOI_TINH}}")
    doc.add_paragraph("4. Số định danh cá nhân: {{SO_DINH_DANH}}")
    doc.add_paragraph("5. Số điện thoại liên hệ: {{DIEN_THOAI}}")
    doc.add_paragraph("6. Email: {{EMAIL}}")
    
    # II. THÔNG TIN CHỦ HỘ
    doc.add_paragraph()
    section2 = doc.add_paragraph()
    section2_run = section2.add_run("II. THÔNG TIN CHỦ HỘ")
    section2_run.bold = True
    
    doc.add_paragraph("1. Họ, chữ đệm và tên chủ hộ: {{CHU_HO}}")
    doc.add_paragraph("2. Mối quan hệ với chủ hộ: {{QUAN_HE_CHU_HO}}")
    doc.add_paragraph("3. Số định danh cá nhân của chủ hộ: {{DINH_DANH_CHU_HO}}")
    
    # III. NỘI DUNG ĐỀ NGHỊ
    doc.add_paragraph()
    section3 = doc.add_paragraph()
    section3_run = section3.add_run("III. NỘI DUNG ĐỀ NGHỊ")
    section3_run.bold = True
    
    doc.add_paragraph("Nội dung đề nghị:")
    doc.add_paragraph("{{NOI_DUNG_DE_NGHI}}")
    
    # IV. Ý KIẾN CỦA CHỦ HỘ
    doc.add_paragraph()
    section4 = doc.add_paragraph()
    section4_run = section4.add_run("IV. Ý KIẾN CỦA CHỦ HỘ")
    section4_run.bold = True
    
    doc.add_paragraph("{{YKIEN_CHU_HO}}")
    
    # V. Ý KIẾN CỦA CHỦ SỞ HỮU
    doc.add_paragraph()
    section5 = doc.add_paragraph()
    section5_run = section5.add_run("V. Ý KIẾN CỦA CHỦ SỞ HỮU CHỖ Ở HỢP PHÁP")
    section5_run.bold = True
    
    doc.add_paragraph("Họ và tên chủ sở hữu: {{CHU_SO_HUU_HO_TEN}}")
    doc.add_paragraph("Số định danh chủ sở hữu: {{CHU_SO_HUU_DINH_DANH}}")
    doc.add_paragraph("Ý kiến: {{YKIEN_CHU_SO_HUU}}")
    
    # VI. Ý KIẾN CỦA CHA, MẸ HOẶC NGƯỜI GIÁM HỘ
    doc.add_paragraph()
    section6 = doc.add_paragraph()
    section6_run = section6.add_run("VI. Ý KIẾN CỦA CHA, MẸ HOẶC NGƯỜI GIÁM HỘ")
    section6_run.bold = True
    
    doc.add_paragraph("Họ và tên cha/mẹ/người giám hộ: {{CHA_ME_HO_TEN}}")
    doc.add_paragraph("Số định danh cha/mẹ/người giám hộ: {{CHA_ME_DINH_DANH}}")
    doc.add_paragraph("Ý kiến: {{YKIEN_CHA_ME}}")
    
    # VII. NGƯỜI KÊ KHAI
    doc.add_paragraph()
    section7 = doc.add_paragraph()
    section7_run = section7.add_run("VII. NGƯỜI KÊ KHAI")
    section7_run.bold = True
    
    doc.add_paragraph("Họ và tên người kê khai: {{NGUOI_KE_KHAI}}")
    
    # Cam kết
    doc.add_paragraph()
    commitment = doc.add_paragraph()
    commitment_run = commitment.add_run("CAM KẾT CỦA NGƯỜI KÊ KHAI")
    commitment_run.bold = True
    
    doc.add_paragraph("Tôi cam kết những thông tin khai trên là đúng sự thật và chịu trách nhiệm trước pháp luật về tính chính xác của những thông tin đã khai.")
    
    # Chữ ký
    doc.add_paragraph()
    doc.add_paragraph("Ngày _____ tháng _____ năm 2024")
    doc.add_paragraph("NGƯỜI KÊ KHAI")
    doc.add_paragraph("(Ký và ghi rõ họ tên)")
    doc.add_paragraph("_________________________")
    
    # Lưu file
    doc.save("ct01_template.docx")
    print("✅ Đã tạo file ct01_template.docx thành công!")

if __name__ == "__main__":
    from docx.shared import Pt
    create_ct01_template() 